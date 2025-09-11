# export_utils.py
import io
from typing import Dict
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


def _flatten_rounds(output: Dict):
    """Return a list of tuples (section_title, data_dict) where each data_dict has
       summaries, insights, questions.
    """
    if "round1" in output:
        return [("Round 1", output["round1"]), ("Round 2", output["round2"])]
    return [("Results", output)]


def build_docx_bytes(topic: str, output: Dict) -> bytes:
    doc = Document()
    styles = doc.styles["Normal"]
    styles.font.name = "Calibri"
    styles.font.size = Pt(11)

    doc.add_heading(f"EulerAI Research Report", 0)
    doc.add_paragraph(f"Topic: {topic}")

    for section_title, data in _flatten_rounds(output):
        doc.add_heading(section_title, level=1)

        doc.add_heading("Summaries", level=2)
        for item in data.get("summaries", []):
            p = doc.add_paragraph()
            p.add_run("Title: ").bold = True
            p.add_run(item["title"])
            doc.add_paragraph(item["summary"])
            if item.get("link"):
                doc.add_paragraph(f"Link: {item['link']}")

        doc.add_heading("Insights", level=2)
        for c, insight in data.get("insights", {}).items():
            doc.add_paragraph(f"[Cluster {c}] {insight}")

        doc.add_heading("Questions", level=2)
        for c, qs in data.get("questions", {}).items():
            doc.add_paragraph(f"[Cluster {c}]")
            doc.add_paragraph(qs)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_pdf_bytes(topic: str, output: Dict) -> bytes:
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4

    def w(text, x=40, y=0):
        # simple writer with wrap
        from reportlab.lib.utils import simpleSplit
        lines = simpleSplit(text, "Helvetica", 10, width - 80)
        yy = y
        for ln in lines:
            c.drawString(x, yy, ln)
            yy -= 14
        return yy

    y = height - 50
    c.setFont("Helvetica-Bold", 14)
    c.drawString(40, y, "EulerAI Research Report")
    y -= 24
    c.setFont("Helvetica", 10)
    y = w(f"Topic: {topic}", y=y)

    for section_title, data in _flatten_rounds(output):
        y -= 18
        c.setFont("Helvetica-Bold", 12)
        c.drawString(40, y, section_title)
        y -= 18
        c.setFont("Helvetica-Bold", 11)
        c.drawString(40, y, "Summaries")
        y -= 18
        c.setFont("Helvetica", 10)
        for item in data.get("summaries", []):
            y = w(f"Title: {item['title']}", y=y)
            y = w(item["summary"], y=y-8)
            if item.get("link"):
                y = w(f"Link: {item['link']}", y=y-8)
            if y < 120:
                c.showPage(); y = height - 50; c.setFont("Helvetica", 10)

        y -= 10
        c.setFont("Helvetica-Bold", 11)
        c.drawString(40, y, "Insights")
        y -= 18; c.setFont("Helvetica", 10)
        for cidx, insight in data.get("insights", {}).items():
            y = w(f"[Cluster {cidx}] {insight}", y=y)
            if y < 120:
                c.showPage(); y = height - 50; c.setFont("Helvetica", 10)

        y -= 10
        c.setFont("Helvetica-Bold", 11)
        c.drawString(40, y, "Questions")
        y -= 18; c.setFont("Helvetica", 10)
        for cidx, qs in data.get("questions", {}).items():
            y = w(f"[Cluster {cidx}]", y=y)
            y = w(qs, y=y-8)
            if y < 120:
                c.showPage(); y = height - 50; c.setFont("Helvetica", 10)

    c.showPage()
    c.save()
    return buf.getvalue()
